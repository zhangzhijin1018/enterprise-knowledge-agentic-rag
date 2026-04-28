"""Document parse API 最小测试。"""

from __future__ import annotations

import pytest
from fastapi.testclient import TestClient

from apps.api.main import app
from core.config import get_settings
from core.repositories.document_chunk_repository import (
    DocumentChunkRepository,
    reset_in_memory_document_chunk_store,
)
from core.repositories.document_repository import DocumentRepository, reset_in_memory_document_store
from core.tools.local.ocr import LocalOCRGateway


@pytest.fixture()
def client(tmp_path, monkeypatch) -> TestClient:
    """创建文档解析 API 测试客户端。"""

    monkeypatch.setenv("LOCAL_UPLOAD_DIR", str(tmp_path / "uploads"))
    get_settings.cache_clear()
    reset_in_memory_document_store()
    reset_in_memory_document_chunk_store()
    with TestClient(app) as test_client:
        yield test_client
    reset_in_memory_document_store()
    reset_in_memory_document_chunk_store()
    get_settings.cache_clear()


def build_auth_headers(user_id: int = 501, username: str = "document_parse_api_user") -> dict[str, str]:
    """构造最小测试认证头。"""

    return {
        "Authorization": "Bearer local-api-token",
        "X-User-Id": str(user_id),
        "X-Username": username,
        "X-Display-Name": username,
        "X-User-Roles": "employee",
        "X-User-Permissions": "document:read,document:write",
        "X-Department-Code": "knowledge-center",
    }


def test_parse_document_generates_parent_child_and_table_chunks(client: TestClient, monkeypatch) -> None:
    """上传文档后应能完成解析、切片并生成父子块和跨页表格关联。"""

    status_transitions: list[str] = []
    original_update_parse_status = DocumentRepository.update_parse_status

    def spy_update_parse_status(self, document_id: str, parse_status: str):
        status_transitions.append(parse_status)
        return original_update_parse_status(self, document_id, parse_status)

    monkeypatch.setattr(DocumentRepository, "update_parse_status", spy_update_parse_status)

    headers = build_auth_headers()
    markdown_content = """# 新能源月度经营报告

## 第一部分 经营概览
本月发电量同比增长，收入稳步提升，整体趋势向好。

表1 发电量统计
| 月份 | 发电量 |
| --- | --- |
| 1月 | 100 |
| 2月 | 120 |
[[PAGE:2]]
| 月份 | 发电量 |
| --- | --- |
| 3月 | 130 |
| 4月 | 140 |

## 第二部分 结论
后续继续提升效率，并优化运行管理制度。
"""

    upload_response = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        files={"file": ("新能源月度经营报告.md", markdown_content.encode("utf-8"), "text/markdown")},
        data={
            "knowledge_base_id": "kb_report_001",
            "business_domain": "report",
            "security_level": "internal",
        },
    )
    document_id = upload_response.json()["data"]["document_id"]

    parse_response = client.post(
        f"/api/v1/documents/{document_id}/parse",
        headers=headers,
    )
    parse_payload = parse_response.json()

    assert parse_response.status_code == 200
    assert parse_payload["data"]["document_id"] == document_id
    assert parse_payload["data"]["parse_status"] == "succeeded"
    assert parse_payload["data"]["chunk_count"] > 0
    assert parse_payload["data"]["parent_chunk_count"] > 0
    assert parse_payload["data"]["child_chunk_count"] > 0
    assert status_transitions == ["processing", "succeeded"]

    detail_response = client.get(f"/api/v1/documents/{document_id}", headers=headers)
    detail_payload = detail_response.json()

    assert detail_response.status_code == 200
    assert detail_payload["data"]["parse_status"] == "succeeded"
    assert detail_payload["data"]["chunk_count"] == parse_payload["data"]["chunk_count"]

    chunk_repository = DocumentChunkRepository(session=None)
    chunks = chunk_repository.list_by_document_id(document_id)

    assert any(item["chunk_type"] == "parent_text" and item["level"] == 1 for item in chunks)
    assert any(
        item["chunk_type"] == "child_text" and item["level"] == 2 and item["parent_chunk_uuid"]
        for item in chunks
    )
    assert any(item["chunk_type"] == "table_parent" for item in chunks)
    assert any(item["chunk_type"] == "table_child" for item in chunks)
    assert any(item["chunk_type"] == "table_summary" for item in chunks)
    assert any(item["metadata"].get("is_cross_page_table") is True for item in chunks)
    assert any(item["metadata"].get("table_group_id") for item in chunks if item["metadata"].get("is_table"))


def test_parse_docx_document_outputs_native_table_blocks(client: TestClient, tmp_path) -> None:
    """docx 解析应直接读取原生段落和表格，而不是依赖 OCR。"""

    from docx import Document as DocxDocument

    headers = build_auth_headers(user_id=601, username="docx_parse_user")
    file_path = tmp_path / "设备检修管理制度.docx"

    document = DocxDocument()
    document.add_heading("设备检修管理制度", level=1)
    document.add_paragraph("第一条 为规范设备检修流程，保障生产安全。")
    document.add_paragraph("表1 检修任务清单")
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "任务"
    table.rows[0].cells[1].text = "负责人"
    table.rows[1].cells[0].text = "检修"
    table.rows[1].cells[1].text = "张三"
    table.rows[2].cells[0].text = "复核"
    table.rows[2].cells[1].text = "李四"
    document.save(file_path)

    with file_path.open("rb") as file_stream:
        upload_response = client.post(
            "/api/v1/documents/upload",
            headers=headers,
            files={
                "file": (
                    file_path.name,
                    file_stream.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={
                "knowledge_base_id": "kb_docx_001",
                "business_domain": "policy",
                "security_level": "internal",
            },
        )

    document_id = upload_response.json()["data"]["document_id"]
    parse_response = client.post(f"/api/v1/documents/{document_id}/parse", headers=headers)

    assert parse_response.status_code == 200
    assert parse_response.json()["data"]["parse_status"] == "succeeded"

    chunks = DocumentChunkRepository(session=None).list_by_document_id(document_id)
    assert any(item["chunk_type"] == "parent_text" for item in chunks)
    assert any(item["chunk_type"] == "table_parent" for item in chunks)
    assert any(item["metadata"].get("table_title") == "表1 检修任务清单" for item in chunks if item["metadata"].get("is_table"))


def test_parse_pdf_document_uses_pdf_text_route(client: TestClient, tmp_path) -> None:
    """文本型 PDF 应走文本抽取路径，并保持 parse 主链路可用。"""

    import fitz

    headers = build_auth_headers(user_id=602, username="pdf_parse_user")
    file_path = tmp_path / "月度经营分析报告.pdf"

    pdf_document = fitz.open()
    page = pdf_document.new_page()
    page.insert_text((72, 72), "第一章 月度经营分析")
    page.insert_text((72, 110), "本月经营情况整体稳定，发电量与收入均稳步增长。")
    pdf_document.save(str(file_path))
    pdf_document.close()

    with file_path.open("rb") as file_stream:
        upload_response = client.post(
            "/api/v1/documents/upload",
            headers=headers,
            files={"file": (file_path.name, file_stream.read(), "application/pdf")},
            data={
                "knowledge_base_id": "kb_pdf_001",
                "business_domain": "report",
                "security_level": "internal",
            },
        )

    document_id = upload_response.json()["data"]["document_id"]
    parse_response = client.post(f"/api/v1/documents/{document_id}/parse", headers=headers)

    assert parse_response.status_code == 200
    assert parse_response.json()["data"]["parse_status"] == "succeeded"
    assert parse_response.json()["data"]["chunk_count"] > 0


def test_parse_image_document_can_trigger_ocr_route(client: TestClient, monkeypatch) -> None:
    """图片文档应能进入 OCR 分支，即使当前测试用 monkeypatch 模拟 OCR 结果。"""

    headers = build_auth_headers(user_id=603, username="ocr_parse_user")
    called = {"ocr": False}

    def fake_extract_structure_from_image(self, image_path, page_no=1, raw_metadata=None):
        called["ocr"] = True
        return [
            {
                "block_type": "ocr_paragraph",
                "text": "设备铭牌信息：风机 1 号，额定功率 2MW。",
                "page_no": page_no,
                "heading_path": ["设备信息"],
                "section_title": "设备信息",
                "clause_no": None,
                "table_data": None,
                "image_ref": str(image_path),
                "raw_metadata": {
                    "parse_mode": "image_ocr_mock",
                    **(raw_metadata or {}),
                },
            }
        ]

    monkeypatch.setattr(LocalOCRGateway, "extract_structure_from_image", fake_extract_structure_from_image)

    upload_response = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        files={"file": ("设备铭牌.png", b"fake-image-bytes", "image/png")},
        data={
            "knowledge_base_id": "kb_img_001",
            "business_domain": "equipment",
            "security_level": "internal",
        },
    )

    document_id = upload_response.json()["data"]["document_id"]
    parse_response = client.post(f"/api/v1/documents/{document_id}/parse", headers=headers)

    assert parse_response.status_code == 200
    assert parse_response.json()["data"]["parse_status"] == "succeeded"
    assert called["ocr"] is True

    chunks = DocumentChunkRepository(session=None).list_by_document_id(document_id)
    assert any(item["chunk_type"] == "parent_text" for item in chunks)
    assert any(item["chunk_type"] == "child_text" for item in chunks)
